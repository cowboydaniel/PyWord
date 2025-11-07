"""
File format handlers for PyWord.

This module provides comprehensive support for importing and exporting documents
in various formats including DOCX, ODT, PDF, RTF, and HTML.
"""

import os
import io
from pathlib import Path
from typing import Optional, Dict, Any, TYPE_CHECKING
from datetime import datetime

if TYPE_CHECKING:
    from .document import Document

# Import format-specific libraries
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from odf.opendocument import load as odf_load, OpenDocumentText
    from odf.text import P, H, Span
    from odf.style import Style, TextProperties, ParagraphProperties
    ODT_AVAILABLE = True
except ImportError:
    ODT_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


class FileFormatHandler:
    """Base class for file format handlers."""

    def __init__(self):
        self.supported = True

    def can_import(self) -> bool:
        """Check if this handler can import files."""
        return False

    def can_export(self) -> bool:
        """Check if this handler can export files."""
        return False

    def import_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Import a document from file.

        Args:
            file_path: Path to the file to import

        Returns:
            Dictionary containing document data or None if failed
        """
        raise NotImplementedError

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """
        Export a document to file.

        Args:
            document: Document to export
            file_path: Path to save the file to

        Returns:
            True if export was successful, False otherwise
        """
        raise NotImplementedError


class DOCXHandler(FileFormatHandler):
    """Handler for Microsoft Word DOCX format."""

    def __init__(self):
        super().__init__()
        self.supported = DOCX_AVAILABLE

    def can_import(self) -> bool:
        return self.supported

    def can_export(self) -> bool:
        return self.supported

    def import_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Import a DOCX document."""
        if not self.supported:
            raise ImportError("python-docx library is not installed")

        try:
            doc = DocxDocument(file_path)

            # Extract text content
            content_parts = []
            for paragraph in doc.paragraphs:
                content_parts.append(paragraph.text)

            content = '\n'.join(content_parts)

            # Extract metadata
            core_props = doc.core_properties

            return {
                'content': content,
                'title': core_props.title or Path(file_path).stem,
                'author': core_props.author or '',
                'created_at': core_props.created.isoformat() if core_props.created else datetime.now().isoformat(),
                'modified_at': core_props.modified.isoformat() if core_props.modified else datetime.now().isoformat(),
                'keywords': core_props.keywords.split(',') if core_props.keywords else [],
                'metadata': {
                    'subject': core_props.subject or '',
                    'category': core_props.category or '',
                    'comments': core_props.comments or ''
                }
            }
        except Exception as e:
            print(f"Error importing DOCX: {e}")
            return None

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """Export a document to DOCX format."""
        if not self.supported:
            raise ImportError("python-docx library is not installed")

        try:
            doc = DocxDocument()

            # Set document properties
            core_props = doc.core_properties
            core_props.title = document.title
            core_props.author = document.author or 'PyWord User'
            if document.keywords:
                core_props.keywords = ','.join(document.keywords)

            # Add content
            if document.content:
                # Split content into paragraphs
                paragraphs = document.content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text)

            # Add sections if available
            for section in document.sections:
                if section.level == 1:
                    doc.add_heading(section.name, level=1)
                elif section.level == 2:
                    doc.add_heading(section.name, level=2)
                else:
                    doc.add_heading(section.name, level=3)

                if section.content:
                    paragraphs = section.content.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text)

            # Save the document
            doc.save(file_path)
            return True

        except Exception as e:
            print(f"Error exporting DOCX: {e}")
            return False


class ODTHandler(FileFormatHandler):
    """Handler for OpenDocument Text (ODT) format."""

    def __init__(self):
        super().__init__()
        self.supported = ODT_AVAILABLE

    def can_import(self) -> bool:
        return self.supported

    def can_export(self) -> bool:
        return self.supported

    def import_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Import an ODT document."""
        if not self.supported:
            raise ImportError("odfpy library is not installed")

        try:
            doc = odf_load(file_path)

            # Extract text content
            content_parts = []
            for element in doc.getElementsByType(P) + doc.getElementsByType(H):
                text_content = self._get_text_content(element)
                if text_content:
                    content_parts.append(text_content)

            content = '\n'.join(content_parts)

            # Extract metadata
            meta = doc.meta
            title = str(meta.getElementsByType(type('title', (), {}))[0]) if meta.getElementsByType(type('title', (), {})) else Path(file_path).stem

            return {
                'content': content,
                'title': title or Path(file_path).stem,
                'author': '',
                'created_at': datetime.now().isoformat(),
                'modified_at': datetime.now().isoformat(),
                'keywords': [],
                'metadata': {}
            }
        except Exception as e:
            print(f"Error importing ODT: {e}")
            return None

    def _get_text_content(self, element) -> str:
        """Extract text content from an ODT element."""
        try:
            # Get all text nodes
            text_parts = []

            def extract_text(node):
                if hasattr(node, 'data'):
                    text_parts.append(node.data)
                if hasattr(node, 'childNodes'):
                    for child in node.childNodes:
                        extract_text(child)

            extract_text(element)
            return ''.join(text_parts)
        except:
            return ''

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """Export a document to ODT format."""
        if not self.supported:
            raise ImportError("odfpy library is not installed")

        try:
            doc = OpenDocumentText()

            # Add content
            if document.content:
                paragraphs = document.content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = P(text=para_text)
                        doc.text.addElement(p)

            # Add sections
            for section in document.sections:
                # Add section heading
                h = H(outlinelevel=section.level, text=section.name)
                doc.text.addElement(h)

                # Add section content
                if section.content:
                    paragraphs = section.content.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            p = P(text=para_text)
                            doc.text.addElement(p)

            # Save the document
            doc.save(file_path)
            return True

        except Exception as e:
            print(f"Error exporting ODT: {e}")
            return False


class PDFHandler(FileFormatHandler):
    """Handler for PDF export."""

    def __init__(self):
        super().__init__()
        self.supported = PDF_AVAILABLE

    def can_import(self) -> bool:
        # PDF import is complex and not supported in this basic implementation
        return False

    def can_export(self) -> bool:
        return self.supported

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """Export a document to PDF format."""
        if not self.supported:
            raise ImportError("reportlab library is not installed")

        try:
            # Create the PDF document
            pdf = SimpleDocTemplate(
                file_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )

            # Container for the 'Flowable' objects
            elements = []

            # Get styles
            styles = getSampleStyleSheet()

            # Add title
            if document.title:
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor='black',
                    spaceAfter=30,
                    alignment=TA_CENTER
                )
                elements.append(Paragraph(document.title, title_style))
                elements.append(Spacer(1, 12))

            # Add metadata
            if document.author:
                author_style = ParagraphStyle(
                    'Author',
                    parent=styles['Normal'],
                    fontSize=12,
                    alignment=TA_CENTER,
                    spaceAfter=12
                )
                elements.append(Paragraph(f"By {document.author}", author_style))
                elements.append(Spacer(1, 12))

            # Add content
            if document.content:
                paragraphs = document.content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = Paragraph(para_text, styles['Normal'])
                        elements.append(p)
                        elements.append(Spacer(1, 12))

            # Add sections
            for section in document.sections:
                # Add section heading
                if section.level == 1:
                    heading = Paragraph(section.name, styles['Heading1'])
                elif section.level == 2:
                    heading = Paragraph(section.name, styles['Heading2'])
                else:
                    heading = Paragraph(section.name, styles['Heading3'])

                elements.append(heading)
                elements.append(Spacer(1, 12))

                # Add section content
                if section.content:
                    paragraphs = section.content.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            p = Paragraph(para_text, styles['Normal'])
                            elements.append(p)
                            elements.append(Spacer(1, 12))

            # Build PDF
            pdf.build(elements)
            return True

        except Exception as e:
            print(f"Error exporting PDF: {e}")
            return False


class RTFHandler(FileFormatHandler):
    """Handler for Rich Text Format (RTF)."""

    def __init__(self):
        super().__init__()
        self.supported = RTF_AVAILABLE

    def can_import(self) -> bool:
        return self.supported

    def can_export(self) -> bool:
        return True  # Basic RTF export is always available

    def import_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Import an RTF document."""
        if not self.supported:
            # Fallback to basic text import
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return {
                    'content': content,
                    'title': Path(file_path).stem,
                    'author': '',
                    'created_at': datetime.now().isoformat(),
                    'modified_at': datetime.now().isoformat(),
                    'keywords': [],
                    'metadata': {}
                }
            except Exception as e:
                print(f"Error importing RTF: {e}")
                return None

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()

            # Convert RTF to plain text
            content = rtf_to_text(rtf_content)

            return {
                'content': content,
                'title': Path(file_path).stem,
                'author': '',
                'created_at': datetime.now().isoformat(),
                'modified_at': datetime.now().isoformat(),
                'keywords': [],
                'metadata': {}
            }
        except Exception as e:
            print(f"Error importing RTF: {e}")
            return None

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """Export a document to RTF format."""
        try:
            # Create a basic RTF document
            rtf_content = self._create_rtf(document)

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(rtf_content)

            return True

        except Exception as e:
            print(f"Error exporting RTF: {e}")
            return False

    def _create_rtf(self, document: 'Document') -> str:
        """Create RTF content from a document."""
        # RTF header
        rtf = r"{\rtf1\ansi\deff0"
        rtf += r"{\fonttbl{\f0 Times New Roman;}}"
        rtf += r"\n"

        # Add title if present
        if document.title:
            rtf += r"\qc\b\fs32 " + self._escape_rtf(document.title) + r"\b0\fs24\par"
            rtf += r"\n"

        # Add author if present
        if document.author:
            rtf += r"\qc\i By " + self._escape_rtf(document.author) + r"\i0\par"
            rtf += r"\n"

        rtf += r"\ql\n"  # Left align

        # Add content
        if document.content:
            paragraphs = document.content.split('\n')
            for para in paragraphs:
                if para.strip():
                    rtf += self._escape_rtf(para) + r"\par"
                    rtf += r"\n"

        # Add sections
        for section in document.sections:
            # Add section heading
            if section.level == 1:
                rtf += r"\b\fs28 " + self._escape_rtf(section.name) + r"\b0\fs24\par"
            elif section.level == 2:
                rtf += r"\b\fs26 " + self._escape_rtf(section.name) + r"\b0\fs24\par"
            else:
                rtf += r"\b " + self._escape_rtf(section.name) + r"\b0\par"

            rtf += r"\n"

            # Add section content
            if section.content:
                paragraphs = section.content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        rtf += self._escape_rtf(para) + r"\par"
                        rtf += r"\n"

        # RTF footer
        rtf += "}"

        return rtf

    def _escape_rtf(self, text: str) -> str:
        """Escape special characters for RTF."""
        text = text.replace('\\', '\\\\')
        text = text.replace('{', r'\{')
        text = text.replace('}', r'\}')
        return text


class HTMLHandler(FileFormatHandler):
    """Handler for HTML format."""

    def __init__(self):
        super().__init__()
        self.supported = HTML_AVAILABLE

    def can_import(self) -> bool:
        return self.supported

    def can_export(self) -> bool:
        return True  # Basic HTML export is always available

    def import_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Import an HTML document."""
        if not self.supported:
            # Fallback to basic text import
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return {
                    'content': content,
                    'title': Path(file_path).stem,
                    'author': '',
                    'created_at': datetime.now().isoformat(),
                    'modified_at': datetime.now().isoformat(),
                    'keywords': [],
                    'metadata': {}
                }
            except Exception as e:
                print(f"Error importing HTML: {e}")
                return None

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract title
            title_tag = soup.find('title')
            title = title_tag.string if title_tag else Path(file_path).stem

            # Extract metadata
            author = ''
            keywords = []

            meta_author = soup.find('meta', attrs={'name': 'author'})
            if meta_author:
                author = meta_author.get('content', '')

            meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
            if meta_keywords:
                keywords = [k.strip() for k in meta_keywords.get('content', '').split(',')]

            # Extract text content
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)

            return {
                'content': content,
                'title': title,
                'author': author,
                'created_at': datetime.now().isoformat(),
                'modified_at': datetime.now().isoformat(),
                'keywords': keywords,
                'metadata': {}
            }
        except Exception as e:
            print(f"Error importing HTML: {e}")
            return None

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """Export a document to HTML format."""
        try:
            html = self._create_html(document)

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html)

            return True

        except Exception as e:
            print(f"Error exporting HTML: {e}")
            return False

    def _create_html(self, document: 'Document') -> str:
        """Create HTML content from a document."""
        html = "<!DOCTYPE html>\n<html>\n<head>\n"
        html += f"    <meta charset=\"UTF-8\">\n"
        html += f"    <title>{self._escape_html(document.title)}</title>\n"

        # Add metadata
        if document.author:
            html += f"    <meta name=\"author\" content=\"{self._escape_html(document.author)}\">\n"

        if document.keywords:
            html += f"    <meta name=\"keywords\" content=\"{self._escape_html(', '.join(document.keywords))}\">\n"

        # Add basic styling
        html += "    <style>\n"
        html += "        body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }\n"
        html += "        h1 { color: #333; }\n"
        html += "        h2 { color: #555; }\n"
        html += "        h3 { color: #777; }\n"
        html += "        p { margin: 10px 0; }\n"
        html += "    </style>\n"
        html += "</head>\n<body>\n"

        # Add title
        if document.title:
            html += f"    <h1>{self._escape_html(document.title)}</h1>\n"

        # Add author
        if document.author:
            html += f"    <p><em>By {self._escape_html(document.author)}</em></p>\n"

        # Add content
        if document.content:
            paragraphs = document.content.split('\n')
            for para in paragraphs:
                if para.strip():
                    html += f"    <p>{self._escape_html(para)}</p>\n"

        # Add sections
        for section in document.sections:
            # Add section heading
            if section.level == 1:
                html += f"    <h2>{self._escape_html(section.name)}</h2>\n"
            elif section.level == 2:
                html += f"    <h3>{self._escape_html(section.name)}</h3>\n"
            else:
                html += f"    <h4>{self._escape_html(section.name)}</h4>\n"

            # Add section content
            if section.content:
                paragraphs = section.content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        html += f"    <p>{self._escape_html(para)}</p>\n"

        html += "</body>\n</html>"

        return html

    def _escape_html(self, text: str) -> str:
        """Escape special characters for HTML."""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&#39;')
        return text


class FileFormatManager:
    """Manager for all file format handlers."""

    def __init__(self):
        self.handlers = {
            'docx': DOCXHandler(),
            'odt': ODTHandler(),
            'pdf': PDFHandler(),
            'rtf': RTFHandler(),
            'html': HTMLHandler(),
            'htm': HTMLHandler(),
        }

    def get_handler(self, extension: str) -> Optional[FileFormatHandler]:
        """Get the appropriate handler for a file extension."""
        ext = extension.lower().lstrip('.')
        return self.handlers.get(ext)

    def can_import(self, extension: str) -> bool:
        """Check if a file format can be imported."""
        handler = self.get_handler(extension)
        return handler is not None and handler.can_import()

    def can_export(self, extension: str) -> bool:
        """Check if a file format can be exported."""
        handler = self.get_handler(extension)
        return handler is not None and handler.can_export()

    def import_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Import a document from file."""
        path = Path(file_path)
        handler = self.get_handler(path.suffix)

        if handler and handler.can_import():
            return handler.import_document(file_path)

        return None

    def export_document(self, document: 'Document', file_path: str) -> bool:
        """Export a document to file."""
        path = Path(file_path)
        handler = self.get_handler(path.suffix)

        if handler and handler.can_export():
            return handler.export_document(document, file_path)

        return False

    def get_supported_import_formats(self) -> list:
        """Get list of supported import formats."""
        formats = []
        for ext, handler in self.handlers.items():
            if handler.can_import():
                formats.append(ext)
        return formats

    def get_supported_export_formats(self) -> list:
        """Get list of supported export formats."""
        formats = []
        for ext, handler in self.handlers.items():
            if handler.can_export():
                formats.append(ext)
        return formats


# Global format manager instance
format_manager = FileFormatManager()
